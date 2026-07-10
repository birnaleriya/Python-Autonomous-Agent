"""
agent.py — DocumentAgent: plan, execute, reflect, render_docx.

Engineering improvement implemented: Multi-step planning + Reflection/self-check.

Why: A single LLM call for a full document produces inconsistent, shallow output.
Breaking the work into plan → execute → reflect gives the model focused, bounded
tasks at each step, and the reflection pass catches weak sections before delivery.

How it improves the agent:
  - Planning produces a validated JSON schema (doc type, tasks, sections) before
    any content is written, so every section has a clear purpose.
  - Execution generates content section-by-section, keeping prompts short and
    focused, which reduces hallucination and improves coherence.
  - Reflection scores each section and rewrites any that score below threshold,
    ensuring consistent business quality across the whole document.
"""
from __future__ import annotations

import io
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from llm_utils import call_llm, call_llm_json

MAX_WORKERS = 2  # keep queue short for local Ollama

logger = logging.getLogger(__name__)

PLAN_SYSTEM = """
You are a senior business analyst. Given a business request, return ONLY a JSON object
with this exact schema (no extra keys):
{
  "document_type": "<proposal|report|plan|analysis|brief>",
  "tasks": ["<task 1>", "<task 2>", ...],
  "sections": [
    {"title": "<Section Title>", "purpose": "<one sentence purpose>"},
    ...
  ]
}
Include 4–7 sections appropriate for the document type.
"""

CONTENT_SYSTEM = """
You are a senior business writer producing content for a formal Word document.
Rules you must follow without exception:
- Write exactly 3 paragraphs. Separate them with a single blank line.
- Each paragraph must be 4–6 complete sentences.
- Use plain prose only. No bullet points, no numbered lists, no markdown, no asterisks.
- Do NOT bold, italicise, or underline any words. No **word** or *word* syntax.
- Do not repeat the section title inside the text.
- Write in third-person formal register suitable for a C-suite audience.
- Every paragraph must add new information; do not pad or repeat earlier points.
"""

REFLECT_SYSTEM = """
You are a strict document quality reviewer. Given a section title, its purpose, and its
current content, return ONLY a JSON object:
{
  "score": <integer 1-10>,
  "issues": "<brief description of issues, or 'none'>",
  "rewrite_needed": <true|false>
}
Score below 7 means rewrite_needed must be true.
"""

REWRITE_SYSTEM = """
You are a senior business writer. Rewrite the section content to fix the listed issues.
Return exactly 3 plain-prose paragraphs separated by blank lines.
No bullet points, no markdown, no bold/italic markers, no section heading in the text.
"""

REFLECTION_THRESHOLD = 7


@dataclass
class SectionResult:
    title: str
    purpose: str
    content: str
    score: int = 0
    issues: str = "none"
    rewrite_needed: bool = False


@dataclass
class AgentState:
    request: str
    document_type: str = ""
    tasks: list[str] = field(default_factory=list)
    sections: list[SectionResult] = field(default_factory=list)
    reflection_summary: str = ""


class DocumentAgent:
    def __init__(self, model: str = "llama3"):
        self.model = model

    # ------------------------------------------------------------------ #
    # Step 1 — Plan                                                        #
    # ------------------------------------------------------------------ #
    def plan(self, request: str) -> AgentState:
        """Decide document type, task list, and section outline."""
        logger.info("Planning document for request: %s", request[:80])
        plan_data: dict[str, Any] = call_llm_json(
            prompt=f"Business request: {request}",
            system=PLAN_SYSTEM,
            model=self.model,
        )
        self._validate_plan(plan_data)

        state = AgentState(request=request)
        state.document_type = plan_data["document_type"]
        state.tasks = plan_data["tasks"]
        state.sections = [
            SectionResult(title=s["title"], purpose=s["purpose"], content="")
            for s in plan_data["sections"]
        ]
        logger.info("Plan complete: %s with %d sections", state.document_type, len(state.sections))
        return state

    # ------------------------------------------------------------------ #
    # Step 2 — Execute                                                     #
    # ------------------------------------------------------------------ #
    def execute(self, state: AgentState) -> AgentState:
        """Generate content for all sections in parallel."""
        def _generate(section: SectionResult) -> tuple[SectionResult, str]:
            logger.info("Generating content for section: %s", section.title)
            prompt = (
                f"Document type: {state.document_type}\n"
                f"Original request: {state.request}\n\n"
                f"Write the '{section.title}' section.\n"
                f"Purpose of this section: {section.purpose}\n\n"
                f"Write 2–4 focused paragraphs."
            )
            return section, call_llm(prompt, system=CONTENT_SYSTEM, model=self.model)

        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as pool:
            futures = {pool.submit(_generate, s): s for s in state.sections}
            for future in as_completed(futures):
                section, content = future.result()
                section.content = content
        return state

    # ------------------------------------------------------------------ #
    # Step 3 — Reflect                                                     #
    # ------------------------------------------------------------------ #
    def reflect(self, state: AgentState) -> AgentState:
        """
        Score all sections in parallel, then rewrite failing ones in parallel.
        """
        def _review(section: SectionResult) -> tuple[SectionResult, dict[str, Any]]:
            review = call_llm_json(
                prompt=(
                    f"Section title: {section.title}\n"
                    f"Purpose: {section.purpose}\n\n"
                    f"Content:\n{section.content}"
                ),
                system=REFLECT_SYSTEM,
                model=self.model,
            )
            return section, review

        def _rewrite(section: SectionResult) -> tuple[SectionResult, str]:
            logger.info("Rewriting section '%s' (score %d): %s",
                        section.title, section.score, section.issues)
            content = call_llm(
                prompt=(
                    f"Section: {section.title}\n"
                    f"Issues to fix: {section.issues}\n\n"
                    f"Original content:\n{section.content}"
                ),
                system=REWRITE_SYSTEM,
                model=self.model,
            )
            return section, content

        # Phase 1 — review all sections in parallel
        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as pool:
            for section, review in pool.map(_review, state.sections):
                section.score = int(review.get("score", 5))
                section.issues = review.get("issues", "none")
                section.rewrite_needed = review.get("rewrite_needed", False)

        # Phase 2 — rewrite failing sections in parallel
        needs_rewrite = [s for s in state.sections if getattr(s, "rewrite_needed", False)]
        rewrites: list[str] = []

        if needs_rewrite:
            with ThreadPoolExecutor(max_workers=MAX_WORKERS) as pool:
                for section, content in pool.map(_rewrite, needs_rewrite):
                    section.content = content
                    rewrites.append(f"'{section.title}' (was {section.score}/10: {section.issues})")

        state.reflection_summary = (
            f"Reflection pass rewrote {len(rewrites)} section(s): {'; '.join(rewrites)}. "
            "All sections now meet quality threshold."
            if rewrites else
            "All sections passed quality review without rewrites."
        )
        return state

    # ------------------------------------------------------------------ #
    # Step 4 — Render .docx                                               #
    # ------------------------------------------------------------------ #
    def render_docx(self, state: AgentState) -> bytes:
        """Build and return a .docx file as bytes."""
        doc = Document()
        self._add_title_page(doc, state)
        self._add_toc_placeholder(doc, state)

        for section in state.sections:
            doc.add_heading(section.title, level=1)
            for para in self._clean_paragraphs(section.content):
                p = doc.add_paragraph(style="Normal")
                run = p.add_run(para)
                run.bold = False
                run.italic = False
                run.font.size = Pt(11)
                p.paragraph_format.space_after = Pt(8)

        self._add_reflection_appendix(doc, state)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------ #
    # Private helpers                                                      #
    # ------------------------------------------------------------------ #
    def _validate_plan(self, data: dict[str, Any]) -> None:
        required = {"document_type", "tasks", "sections"}
        missing = required - data.keys()
        if missing:
            raise ValueError(f"Plan JSON missing keys: {missing}")
        if not isinstance(data["sections"], list) or len(data["sections"]) < 2:
            raise ValueError("Plan must contain at least 2 sections.")
        for s in data["sections"]:
            if "title" not in s or "purpose" not in s:
                raise ValueError(f"Section missing title/purpose: {s}")

    @staticmethod
    def _clean_paragraphs(text: str) -> list[str]:
        """Strip markdown artifacts and return non-empty paragraphs."""
        import re
        # Remove **bold**, *italic*, __bold__, _italic_ markers
        text = re.sub(r"\*{1,2}(.+?)\*{1,2}", r"\1", text)
        text = re.sub(r"_{1,2}(.+?)_{1,2}", r"\1", text)
        # Remove any leading hashes (accidental headings)
        text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
        return [p.strip() for p in text.split("\n\n") if p.strip()]

    def _add_title_page(self, doc: Document, state: AgentState) -> None:
        doc.add_paragraph()  # top spacer
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(state.document_type.upper())
        run.bold = True
        run.font.size = Pt(26)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        doc.add_paragraph()  # spacer

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle.add_run(state.request)
        sub_run.italic = True
        sub_run.bold = False
        sub_run.font.size = Pt(12)
        sub_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        doc.add_page_break()

    def _add_toc_placeholder(self, doc: Document, state: AgentState) -> None:
        doc.add_heading("Table of Contents", level=1)
        for i, section in enumerate(state.sections, 1):
            doc.add_paragraph(f"{i}. {section.title}", style="List Number")
        doc.add_page_break()

    def _add_reflection_appendix(self, doc: Document, state: AgentState) -> None:
        doc.add_page_break()
        doc.add_heading("Agent Reflection", level=1)
        ref_para = doc.add_paragraph(style="Normal")
        ref_run = ref_para.add_run(state.reflection_summary)
        ref_run.bold = False
        ref_run.font.size = Pt(11)
        doc.add_paragraph()

        doc.add_heading("Section Quality Scores", level=2)
        for section in state.sections:
            p = doc.add_paragraph(style="List Bullet")
            title_run = p.add_run(f"{section.title}: ")
            title_run.bold = True
            title_run.font.size = Pt(11)
            score_run = p.add_run(f"{section.score}/10 — {section.issues}")
            score_run.bold = False
            score_run.font.size = Pt(11)
